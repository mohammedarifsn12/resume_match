import streamlit as st
import docx2txt
from PyPDF2 import PdfReader
from langdetect import detect
import io
import re
from docx import Document

# Set up Streamlit app
st.set_page_config(page_title="SkillSync", page_icon="🛠️", layout="wide")
st.title("🛠️ SkillSync: Your Resume & Career Companion 💼")

# Mapped terms for short forms
KEYWORD_MAPPINGS = {
    "Machine Learning": ["ML", "machine learning", "machine learnin", "m/l", "machine-learning", "ml algorithms"],
    "Search Engine Optimization": ["SEO", "search optimization", "search engine optimization", "seo strategy", "seo techniques", "site optimization"],
    "Data Analysis": ["data analytics", "analysis", "data analysis", "data interpretation", "analyzing data", "data insights"],
    "JavaScript": ["JS", "javascript", "java script", "js frameworks", "js programming", "javascript coding"],
    "Natural Language Processing": ["NLP", "text processing", "natural lang proc", "nlp models", "language processing", "natural-language-processing"],
    "Penetration Testing": ["Pentest", "security testing", "penetration test", "vulnerability assessment", "pentesting", "ethical hacking"],
    "Python": ["py", "python programming", "python scripting", "python lang", "python code", "pythonic"],
    "SQL": ["structured query language", "sql queries", "database querying", "sql scripting", "sql commands", "relational database"],
    "Deep Learning": ["DL", "deep-learning", "deep neural networks", "deep nn", "dl algorithms", "deep neural net"],
    "Data Visualization": ["data viz", "data visualization", "charts", "graphs", "dashboards", "visualizing data"],
    "Project Management": ["PM", "project mgr", "project handling", "proj mgmt", "project coordination", "project leadership"],
    "Agile": ["scrum", "agile methodology", "agile framework", "agile processes", "agile project management", "agile software development"],
    "Cloud Computing": ["cloud infra", "cloud infrastructure", "cloud services", "aws", "azure", "gcp"],
    "UI/UX Design": ["user interface design", "ux design", "user experience", "ui/ux", "interface design", "ux/ui"],
    "Mobile Development": ["mobile dev", "app dev", "mobile app development", "mobile programming", "app creation", "native apps"],
    "Cybersecurity": ["network security", "cyber security", "cybersec", "information security", "infosec", "cyber protection"],
    "SEO": ["seo", "search engine opt", "website optimization", "seo marketing", "seo campaigns", "seo strategies"],
    "Artificial Intelligence": ["AI", "artificial intelligence", "ai models", "ai techniques", "machine intelligence", "artificial-intelligence"],
    "Data Engineering": ["data engg", "data pipelines", "data engineering", "etl processes", "data processing", "data integration"],
    "Big Data": ["bigdata", "big-data", "data lakes", "hadoop", "spark", "big data technologies"],
    "DevOps": ["devops", "ci/cd", "continuous integration", "continuous delivery", "docker", "kubernetes"],
    "Blockchain": ["blockchain", "distributed ledger", "cryptocurrency", "ethereum", "smart contracts", "dapps"],
    "Quality Assurance": ["qa", "testing", "quality testing", "automation testing", "manual testing", "bug tracking"],
    "Data Governance": ["data policies", "data compliance", "data stewardship", "data privacy", "data management"],
    "Leadership": ["team management", "mentorship", "decision making", "vision planning", "strategy development"],
    "Communication": ["public speaking", "presentation skills", "writing", "verbal communication", "interpersonal skills"],
    "Teamwork": ["collaboration", "team player", "cross-functional collaboration", "working with teams", "team coordination"],
    "Problem Solving": ["critical thinking", "problem solving", "analytical thinking", "troubleshooting", "solution-oriented"],
    "Adaptability": ["flexibility", "adaptability", "open to change", "quick learner", "resilience", "work under pressure"],
    "Time Management": ["time management", "prioritization", "task management", "deadline driven", "multitasking", "efficiency"],
    "Creativity": ["innovation", "creative thinking", "idea generation", "design thinking", "out-of-the-box thinking"],
    "Microsoft Word": ["microsoft word", "word processing", "document formatting", "creating reports", "word documents"],
    "Microsoft Excel": ["microsoft excel", "data analysis", "spreadsheet", "pivot tables", "formulas", "excel macros"],
    "Microsoft PowerPoint": ["microsoft powerpoint", "presentation slides", "ppt", "slide design", "powerpoint presentations"],
    "Microsoft Outlook": ["microsoft outlook", "email communication", "outlook calendar", "meeting scheduling", "email management"],
    "Microsoft Access": ["microsoft access", "database management", "queries", "relational databases", "data entry"],
    "LMS Administration": ["lms", "learning management system", "training support"],
    "Microsoft Office": ["word", "excel", "powerpoint"],
    "Data Science": ["data science", "data scientist", "data-driven insights", "data analysis models", "data research"],
    "Graphic Design": ["graphic design", "photoshop", "illustrator", "adobe creative suite", "design concepts"],
    "Web Development": ["web dev", "website development", "frontend coding", "backend development", "web programming"],
    "Internet of Things": ["IoT", "smart devices", "internet of things", "iot technologies", "connected systems"],
    "Robotic Process Automation": ["RPA", "robotic automation", "automation bots", "process automation", "rpa tools"],
    "Game Development": ["game dev", "game programming", "unity", "unreal engine", "game design"],
    "Augmented Reality": ["AR", "augmented reality", "mixed reality", "ar design", "ar applications"],
    "Virtual Reality": ["VR", "virtual reality", "vr development", "virtual experiences", "immersive tech"],
    "Business Analysis": ["business analytics", "requirements gathering", "business needs analysis", "BA role"],
    "Operations Management": ["ops management", "operations efficiency", "supply chain", "logistics", "ops planning"],
    "Technical Writing": ["tech writing", "documentation", "manual writing", "technical authoring", "content creation"],
    "Digital Transformation": ["digital strategy", "digital transformation initiatives", "business digitization"],
    "Human Resources": ["HR", "recruitment", "talent management", "employee relations", "hr processes"],
    "Content Marketing": ["content strategy", "blogging", "content creation", "social media posts", "content marketing"],
    "Email Marketing": ["email campaigns", "email strategy", "drip campaigns", "email funnels", "mail automation"],
    "IT Support": ["tech support", "it troubleshooting", "desktop support", "helpdesk", "technical assistance"],
    "Social Media Management": ["social platforms", "scheduling posts", "content curation", "social engagement"],
    "E-commerce Management": ["online store management", "product listing", "inventory management", "ecommerce analytics"],
    "Financial Analysis": ["financial modeling", "investment analysis", "risk assessment", "financial reporting"],
    "Supply Chain Management": ["logistics management", "warehouse optimization", "supply chain processes"],
    "Customer Service": ["customer support", "client relations", "customer care", "support resolution"],
    "Cloud Architecture": ["cloud solutions", "cloud architecture design", "cloud scaling"],
    "Network Administration": ["network setup", "IT infrastructure", "network security"],
    "Data Warehousing": ["data lake", "ETL tools", "database optimization", "warehouse solutions"],
    "Database Administration": ["db admin", "database servers", "SQL tuning", "database backup"],
    "AI Ethics": ["AI fairness", "algorithmic transparency", "responsible AI"],
    "UI Design": ["interface design principles", "prototyping", "user testing"],
    "UX Research": ["user behavior analysis", "journey mapping", "usability testing"],
    "Mobile UI Design": ["responsive mobile design", "mobile-first", "touch interface design"],
    "Embedded Systems": ["microcontrollers", "embedded programming", "real-time systems"],
    "Digital Forensics": ["cyber investigations", "digital evidence", "incident analysis"],
    "Risk Management": ["risk assessment", "mitigation strategies", "compliance risks"],
    "DevSecOps": ["secure devops", "security pipelines", "devops compliance"],
    "3D Modeling": ["3d rendering", "model design", "maya", "blender"],
    "Video Editing": ["video post-production", "premiere pro", "final cut pro", "video content creation"],
    "Audio Engineering": ["sound mixing", "audio production", "studio recording"],
    "Renewable Energy": ["solar tech", "green energy", "wind turbine tech", "sustainability projects"],
    "Event Management": ["event planning", "conference logistics", "event marketing"],
    "Legal Compliance": ["regulatory compliance", "legal frameworks", "policy adherence"],
    "Data Security": ["data protection", "information security measures", "data encryption"],
    "Hardware Design": ["PCB design", "hardware testing", "circuit simulation"],
    "Quantum Computing": ["qubits", "quantum algorithms", "quantum programming"],
    "Mathematical Modeling": ["numerical simulations", "optimization problems", "mathematical proofs"],
    "ERP Systems": ["enterprise resource planning", "SAP ERP", "oracle ERP", "workflow automation"],
    "CRM Systems": ["salesforce CRM", "customer relations tools", "CRM integrations"],
    "API Development": ["REST API", "SOAP API", "API integrations"],
    "Ethical Hacking": ["penetration testing certifications", "vulnerability exploitation"],
    "Mobile Security": ["secure apps", "mobile device management", "msecure practices"],
    "System Design": ["scalable architectures", "design diagrams", "system workflows"],
    "SaaS Solutions": ["subscription platforms", "software as a service", "SaaS optimization"],
    "Biometric Technology": ["facial recognition", "fingerprint systems", "biometric auth"],
    "Data Ethics": ["data anonymization", "privacy ethics", "ethical data use"],
    "Personal Branding": ["linkedin optimization", "professional profiles", "online visibility"],
    "Cryptography": ["data encryption techniques", "secure communications", "cryptanalysis"],
    "Drone Technology": ["aerial robotics", "drone operations", "unmanned vehicles"],
    "E-Learning": ["course creation", "online teaching", "virtual classrooms"],
    "Cross-Platform Development": ["flutter", "react native", "cross-platform apps"],
    "Data Archiving": ["long-term storage", "archive solutions", "data lifecycle management"],
    "Remote Sensing": ["satellite imagery", "geospatial analysis", "remote sensing tech"],
    "Biomedical Engineering": ["bioinformatics", "medical imaging", "healthcare innovations"],
    "Climate Modeling": ["weather predictions", "climate simulations", "earth systems analysis"],
    "Robotics": ["autonomous robots", "mechanical systems", "robot control software"],
    "Crowdfunding": ["fundraising campaigns", "kickstarter", "indiegogo marketing"],
    "Negotiation": ["contract negotiation", "deal closure", "conflict resolution"],
    "Organizational Behavior": ["team dynamics", "corporate culture", "leadership impact"],
    "Storyboarding": ["narrative planning", "animation storyboards", "visual scripts"],
    "Design Thinking": ["empathize ideate prototype", "innovation strategy", "design processes"],
    "Health Informatics": ["electronic health records", "medical data systems", "patient management tech"],
    "Bioinformatics": ["genomic data", "protein analysis", "bioinfo tools"],
    "Agricultural Technology": ["precision farming", "agribots", "soil data analysis"],
    "Digital Wallets": ["payment systems", "mobile wallets", "digital payments"],
    "Smart Cities": ["urban IoT", "smart grids", "smart mobility"],
    "Cryptocurrency": ["crypto trading", "bitcoin", "blockchain finance"],
    "Creative Writing": ["fiction writing", "scriptwriting", "creative content"],
    "Public Relations": ["media relations", "press releases", "PR campaigns"],
    "Space Technologies": ["satellite systems", "space exploration", "aerospace engineering"],
    "3D Printing": ["additive manufacturing", "printing tech", "material science"],
    "Open Source Contributions": ["github projects", "open-source codebases", "community software"],
    "Statistical Modeling": ["regression models", "time series analysis", "statistical inference"],
    "Ethics and Compliance": ["ethical policies", "corporate compliance", "values alignment"],
    "Web Security": ["anti-hacking", "secure protocols", "site hardening"],
    "Digital Branding": ["brand management", "online identity", "digital campaigns"],
    "Digital Twins": ["virtual replicas", "simulation models", "digital mapping"],
    "UX Strategy": ["user-centered design", "experience optimization", "strategy alignment"],
    "AI in Healthcare": ["predictive healthcare", "ai diagnostics", "medtech innovation"],
    "Bioengineering": ["biotech", "bioengineering processes", "biological systems", "biomaterials", "biosensors"],
    "Knowledge Management": ["organizational knowledge", "KM tools", "knowledge sharing", "knowledge retention"],
    "Digital Advertising": ["online ads", "ad targeting", "ppc campaigns", "digital ad strategies"],
    "Industrial Automation": ["factory automation", "industrial robots", "process control", "PLC programming"],
    "Sustainability": ["eco-friendly practices", "green technologies", "sustainability metrics", "carbon footprint reduction"],
    "Energy Management": ["energy optimization", "renewable resources", "power distribution", "smart grids"],
    "Customer Experience": ["CX", "customer journey mapping", "user experience improvement", "customer touchpoints"],
    "Market Research": ["consumer insights", "competitive analysis", "market trends", "survey methodologies"],
    "User Research": ["user feedback", "persona development", "UX insights", "research methodologies"],
    "Predictive Analytics": ["trend forecasting", "data predictions", "predictive modeling", "forecasting tools"],
    "Knowledge Graphs": ["semantic graphs", "data relationships", "ontology building", "linked data"],
    "Edge Computing": ["fog computing", "distributed systems", "local data processing", "edge devices"],
    "Hybrid Cloud": ["cloud combination", "public-private cloud", "hybrid IT solutions", "hybrid infrastructure"],
    "Open Data": ["data democratization", "public datasets", "open access information", "data transparency"],
    "Social Impact Analysis": ["community initiatives", "impact metrics", "social project evaluations"],
    "Animation": ["2D animation", "3D animation", "motion graphics", "animation workflows"],
    "Corporate Training": ["employee development", "training programs", "workshop facilitation"],
    "Smart Devices": ["IoT gadgets", "connected tech", "home automation", "smart system integration"],
    "Health Monitoring": ["wearable tech", "health sensors", "patient monitoring", "fitness trackers"],
    "Server Administration": ["server maintenance", "hosting management", "system performance", "server-side support"],
    "Environmental Science": ["climate studies", "biodiversity research", "eco systems", "sustainability science"],
    "Ethnographic Studies": ["cultural insights", "field research", "human behavior analysis"],
    "Legal Tech": ["contract automation", "legal research tools", "case management software"],
    "Quantum Algorithms": ["quantum programming", "shor's algorithm", "quantum computing models"],
    "Crowd Management": ["public safety", "event security", "crowd analytics", "large-scale coordination"],
    "AI-Powered Chatbots": ["conversational AI", "customer service bots", "virtual assistants"],
    "Media Production": ["content editing", "studio operations", "multimedia creation"],
    "Climate Adaptation": ["resilience planning", "climate risk assessment", "adaptive measures"],
    "E-Waste Management": ["electronics recycling", "waste minimization", "eco disposal"],
    "RegTech": ["regulatory tech", "compliance automation", "legal risk management"],
    "Product Lifecycle Management": ["PLM", "design-to-market", "lifecycle optimization"],
    "Agritech": ["farming tech", "smart agriculture", "crop monitoring", "farm optimization"],
    "Behavioral Science": ["human behavior patterns", "decision analysis", "behavior modeling"],
    "Language Translation": ["linguistic services", "automated translation", "multilingual communication"],
    "Digital Twins": ["real-time simulation", "virtual replicas", "digital modeling"],
    "Emotional Intelligence": ["EQ", "empathy building", "emotional awareness"],
    "Smart Wearables": ["wearable tech", "smart health devices", "fitness bands"],
    "Space Exploration": ["orbital systems", "planetary research", "aerospace innovations"],
    "Data Privacy": ["GDPR compliance", "privacy by design", "data protection laws"],
    "Speech Processing": ["voice recognition", "speech-to-text", "acoustic analysis"],
    "Text Summarization": ["extractive summarization", "abstractive summarization", "text compression"],
    "Knowledge Representation": ["logical reasoning", "ontology creation", "semantic frameworks"],
    "Bioinformatics Algorithms": ["genomic sequence", "phylogenetic trees", "biological computation"],
    "Robotic Surgery": ["automated procedures", "surgical robotics", "minimally invasive tech"],
    "Digital Storytelling": ["interactive narratives", "multimedia stories", "virtual storytelling"],
    "IoT Security": ["device protection", "IoT encryption", "network integrity"],
    "API Security": ["token management", "secure endpoints", "API gateway protection"],
    "Low-Code Development": ["drag-and-drop coding", "visual programming", "rapid app development"],
    "Game Mechanics": ["reward systems", "level design", "player interaction"],
    "Data Compression": ["file size reduction", "lossless compression", "encoding techniques"],
    "Image Processing": ["pixel manipulation", "filter design", "image segmentation"],
    "Energy Storage": ["battery tech", "grid storage", "power cells"],
    "Vehicle Telematics": ["fleet management", "vehicle tracking", "telematic systems"],
    "Performance Optimization": ["process improvement", "speed enhancement", "resource tuning"],
    "Employee Engagement": ["workforce satisfaction", "motivation strategies", "engagement metrics"],
    "Recruitment Technology": ["AI hiring tools", "resume screening", "job matching"],
    "Fraud Detection": ["anomaly spotting", "risk algorithms", "transaction monitoring"],
    "Crisis Management": ["contingency planning", "risk mitigation", "emergency response"],
    "Data Provenance": ["source tracking", "data lineage", "traceability"],
    "Algorithm Development": ["problem-solving algorithms", "optimization techniques"],
    "Oceanography": ["marine studies", "ocean modeling", "underwater exploration"],
    "Wearable Robotics": ["exoskeleton tech", "assistive devices", "robotic wearables"],
    "Ethical HCI": ["user ethics", "human-centered computing", "design responsibility"],
    "Remote Collaboration": ["virtual teams", "online brainstorming", "remote work tools"],
    "Scripting Languages": ["bash scripting", "perl programming", "automation scripts"],
    "Mobile AR": ["augmented mobile apps", "AR gaming", "AR filters"],
    "Assistive Technology": ["disability tech", "assistive devices", "accessible solutions"],
    "Cross-Cultural Training": ["cultural sensitivity", "global business training", "cross-cultural dynamics"],
    "Data Augmentation": ["synthetic data", "training data enhancement", "augmentation techniques"],
    "Microservices Architecture": ["service-oriented architecture", "SOA", "microservices development", "containerized apps"],
    "Renewable Energy": ["solar energy", "wind power", "green energy", "sustainable energy systems"],
    "Digital Marketing Analytics": ["campaign metrics", "marketing KPIs", "ROI tracking", "conversion analysis"],
    "Computer Vision": ["image recognition", "object detection", "CV algorithms", "vision systems"],
    "Human-Computer Interaction": ["HCI design", "user interfaces", "interaction design", "user experience studies"],
    "Gene Editing": ["CRISPR", "genetic modification", "biotech engineering", "gene therapy"],
    "Video Production": ["film editing", "video scripting", "cinematography", "media editing"],
    "Gamification": ["engagement strategies", "game-like experiences", "behavioral incentives"],
    "Supply Chain Optimization": ["logistics management", "inventory planning", "SCM", "demand forecasting"],
    "e-Learning Development": ["online course creation", "instructional design", "e-learning platforms"],
    "Data Lakes": ["big data storage", "data lake architecture", "unstructured data handling"],
    "Digital Twins for IoT": ["virtual IoT devices", "real-time modeling", "digital representation"],
    "5G Technology": ["next-gen networks", "5G connectivity", "ultra-fast broadband", "wireless tech"],
    "Autonomous Vehicles": ["self-driving cars", "vehicle autonomy", "ADAS", "driverless systems"],
    "Kubernetes Orchestration": ["K8s", "container management", "cluster orchestration"],
    "Data Monetization": ["data-driven revenue", "data as a product", "commercializing data"],
    "Neuro-Linguistic Programming": ["NLP therapy", "mindset transformation", "communication techniques"],
    "Chatbot Development": ["virtual agents", "AI conversations", "interactive bots"],
    "Customer Retention": ["loyalty programs", "churn reduction", "client engagement strategies"],
    "Knowledge-Based Systems": ["expert systems", "decision support", "rule-based AI"],
    "Wearable Medical Devices": ["health trackers", "biometric monitoring", "wearable diagnostics"],
    "Hybrid Work Models": ["remote and onsite", "flexible work", "hybrid workforce strategies"],
    "AI in Healthcare": ["medical imaging AI", "predictive health analytics", "digital diagnosis"],
    "Smart Grids": ["intelligent energy networks", "grid optimization", "smart energy systems"],
    "Automated Quality Control": ["QA automation", "smart testing systems", "machine-driven QC"],
    "Personal Finance Management": ["budget tracking", "expense management", "financial planning tools"],
    "Voice Assistants": ["virtual assistants", "voice AI", "voice-activated tech"],
    "Predictive Maintenance": ["equipment monitoring", "failure prediction", "maintenance analytics"],
    "Generative AI": ["creative AI", "content generation", "AI-generated art"],
    "Sentiment Analysis": ["opinion mining", "customer sentiment", "text-based emotion analysis"],
    "Natural Disaster Management": ["disaster prediction", "emergency systems", "crisis response tech"],
    "Information Retrieval": ["search algorithms", "data fetching", "indexing systems"],
    "Cloud Security": ["data protection in cloud", "cloud access control", "cloud encryption"],
    "AI Ethics": ["responsible AI", "bias mitigation", "ethical decision-making in AI"],
    "Social Media Management": ["content scheduling", "social media campaigns", "SM analytics"],
    "IoT Analytics": ["device data analysis", "IoT insights", "connected device stats"],
    "Synthetic Biology": ["bioengineered systems", "synthetic organisms", "genome synthesis"],
    "Predictive Text": ["autocomplete systems", "AI typing", "text prediction"],
    "AI-Powered Personalization": ["custom user experiences", "personalized recommendations", "AI-driven customization"],
    "Data Labeling": ["annotation tools", "data tagging", "training data prep"],
    "Home Automation": ["smart home tech", "connected living", "home IoT devices"],
    "Advanced Robotics": ["robotics AI", "intelligent automation", "robot dynamics"],
    "E-Commerce Optimization": ["online sales strategies", "cart abandonment reduction", "conversion rate optimization"],
    "Workflow Automation": ["process automation", "task orchestration", "workflow tools"],
    "Design Systems": ["UI frameworks", "design libraries", "component-based design"],
    "Streaming Media": ["live streaming", "OTT platforms", "video on demand"],
    "Ethical Hacking": ["penetration testing", "network defense", "ethical cybersecurity"],
    "Robotics Process Automation": ["RPA bots", "task automation", "workflow efficiency"],
    "Digital Watermarking": ["content protection", "media ownership", "digital rights"],
    "Crowdsourcing": ["distributed problem-solving", "crowd collaboration", "open innovation"],
    "Virtual Events": ["online conferencing", "digital summits", "virtual meetups"],
    "Real-Time Analytics": ["instant insights", "live data processing", "real-time dashboards"],
    "Mobile Wallets": ["digital payments", "contactless transactions", "e-wallets"],
    "Affective Computing": ["emotion AI", "sentiment-aware systems", "human-like interactions"],
    "Cyber Threat Intelligence": ["threat hunting", "cyber risk analysis", "malware detection"],
    "Zero Trust Security": ["access verification", "identity-first security", "ZT architecture"],
    "Customer Onboarding": ["client activation", "welcome programs", "onboarding strategies"],
    "Biometrics": ["fingerprint recognition", "facial authentication", "biometric security"],
    "SaaS Development": ["cloud applications", "software as a service", "subscription models"],
    "AI-Powered Creativity": ["creative assistants", "AI design tools", "generative design"],
    "Digital Therapeutics": ["health apps", "digital care programs", "behavioral health tech"],
    "Account Management": ["key account management", "client account handling", "account relationships"],
    "Business Development": ["BD", "sales growth", "new client acquisition", "expansion strategies"],
    "Sales Enablement": ["sales support", "enablement tools", "sales productivity"],
    "Client Relationship Management": ["CRM", "customer relationships", "client interactions", "relationship building"],
    "Lead Generation": ["prospecting", "lead identification", "new business leads"],
    "Sales Strategy": ["sales planning", "strategic selling", "revenue growth plans"],
    "Customer Retention": ["client loyalty", "retention programs", "renewals"],
    "Pipeline Management": ["sales pipeline", "opportunity tracking", "deal progression"],
    "Cross-Selling": ["additional sales", "product upsell", "client upsell opportunities"],
    "Client Onboarding": ["welcome process", "initial engagement", "onboarding support"],
    "B2B Sales": ["business-to-business sales", "corporate sales", "enterprise deals"],
    "B2C Sales": ["retail sales", "consumer engagement", "direct selling"],
    "Sales Forecasting": ["revenue prediction", "sales trends", "performance estimation"],
    "Cold Calling": ["outbound calls", "prospecting calls", "lead cold outreach"],
    "Customer Feedback Analysis": ["client surveys", "customer insights", "feedback review"],
    "Sales Presentations": ["pitch decks", "client proposals", "sales demonstrations"],
    "Negotiation Skills": ["deal negotiation", "contract finalization", "sales closures"],
    "Sales Reporting": ["performance metrics", "sales data analysis", "revenue reports"],
    "Upselling Strategies": ["upsell techniques", "premium sales", "product upgrades"],
    "Customer Complaint Handling": ["issue resolution", "escalation management", "client grievances"],
    "Retail Management": ["store sales", "in-store promotions", "retail strategy"],
    "Territory Management": ["regional sales", "territory planning", "area coverage"],
    "Market Research for Sales": ["competitor analysis", "market trends", "industry insights"],
    "Sales Incentives Planning": ["commission structures", "bonus plans", "sales motivation"],
    "Inside Sales": ["remote sales", "virtual client interactions", "phone-based selling"],
    "Field Sales": ["door-to-door sales", "on-site client visits", "outdoor sales"],
    "Channel Sales": ["partner networks", "distributor management", "reseller programs"],
    "After-Sales Service": ["customer support", "service follow-up", "post-sales care"],
    "Sales Training": ["team coaching", "sales workshops", "skill development"],
    "Customer Advocacy": ["client champions", "customer reference programs", "client success stories"],
    "Key Performance Indicators (KPIs)": ["sales KPIs", "performance tracking", "goal achievement"],
    "Proposal Writing": ["client proposals", "business bids", "sales documentation"],
    "Loyalty Programs": ["customer rewards", "membership programs", "repeat business incentives"],
    "CRM Software": ["customer management tools", "salesforce", "hubspot", "zoho CRM"],
    "Product Demos": ["live demonstrations", "feature presentations", "trial sessions"],
    "Contract Negotiation": ["deal terms", "agreement finalization", "contract drafting"],
    "Sales Funnel Optimization": ["conversion rate", "sales journey", "lead nurturing"],
    "Event Sales": ["exhibition sales", "trade show leads", "event-based promotions"],
    "Client Escalation Handling": ["issue escalation", "critical client cases", "problem resolution"],
    "Sales Automation Tools": ["sales tech", "pipeline automation", "sales software"],
    "Consultative Selling": ["problem-solving sales", "solution-driven selling", "consultative approaches"],
    "Customer Journey Mapping": ["client lifecycle", "user journey", "customer path analysis"],
    "Proposal Negotiation": ["proposal approval", "deal crafting", "terms negotiation"],
    "Value-Based Selling": ["ROI selling", "customer value creation", "benefit-driven sales"],
    "Client Retention Strategies": ["long-term client relationships", "loyalty building", "repeat business"],
    "Social Selling": ["sales via social media", "LinkedIn prospecting", "social engagement"],
    "Client Portfolio Management": ["client categorization", "portfolio optimization", "high-value clients"],
    "Order Fulfillment": ["order tracking", "delivery coordination", "sales order execution"],
    "Brand Ambassador Programs": ["sales advocacy", "client representation", "customer promoters"]
}

CATEGORIES = {
    "Data Science": {
        "Skills": ["Python", "Machine Learning", "Data Analysis", "SQL", "Statistics", "Data Visualization", "Big Data"],
        "Action Verbs": ["Analyzed", "Optimized", "Predicted", "Validated", "Integrated"],
        "Quantifiers": ["90%", "Millions", "Billions", "10x", "Hundreds"]
    },
    "Web Development": {
        "Skills": ["HTML", "CSS", "JavaScript", "React", "Node.js", "TypeScript", "Webpack", "Bootstrap"],
        "Action Verbs": ["Developed", "Implemented", "Debugged", "Optimized", "Customized"],
        "Quantifiers": ["500+", "25%", "Thousands", "Multiple", "15x"]
    },
    "Mobile Development": {
        "Skills": ["Kotlin", "Swift", "React Native", "Flutter", "Java", "Objective-C", "Firebase"],
        "Action Verbs": ["Built", "Tested", "Designed", "Streamlined", "Published"],
        "Quantifiers": ["100+", "5x", "Hundreds", "Thousands", "20+"]
    },
    "Cybersecurity": {
        "Skills": ["Network Security", "Penetration Testing", "Risk Assessment", "Cryptography", "SOC Operations", "Firewall Management"],
        "Action Verbs": ["Protected", "Secured", "Mitigated", "Evaluated", "Prevented"],
        "Quantifiers": ["95%", "Thousands", "Zero", "Multiple", "Dozens"]
    },
    "Digital Marketing": {
        "Skills": ["SEO", "Google Analytics", "Content Marketing", "PPC", "Social Media Strategy", "Email Campaigns"],
        "Action Verbs": ["Improved", "Increased", "Promoted", "Managed", "Enhanced"],
        "Quantifiers": ["300%", "Double", "50%", "Thousands", "Hundreds"]
    },
    "Product Management": {
        "Skills": ["Agile", "Scrum", "Roadmap Planning", "Stakeholder Management", "Market Research", "User Story Writing"],
        "Action Verbs": ["Led", "Directed", "Coordinated", "Prioritized", "Facilitated"],
        "Quantifiers": ["Multiple", "5+", "Millions", "Zero", "20%"]
    },
    "UI/UX Design": {
        "Skills": ["Figma", "Sketch", "Wireframing", "Prototyping", "User Research", "Adobe XD", "Accessibility Design"],
        "Action Verbs": ["Designed", "Prototyped", "Conducted", "Iterated", "Enhanced"],
        "Quantifiers": ["Dozens", "Thousands", "Multiple", "10+", "15%"]
    },
    "Cloud Computing": {
        "Skills": ["AWS", "Azure", "Google Cloud", "DevOps", "Terraform", "Kubernetes", "Cloud Security"],
        "Action Verbs": ["Deployed", "Scaled", "Automated", "Orchestrated", "Provisioned"],
        "Quantifiers": ["Millions", "Hundreds", "Zero", "5x", "50%"]
    },
    "AI/ML Engineering": {
        "Skills": ["Deep Learning", "NLP", "TensorFlow", "PyTorch", "Computer Vision", "Reinforcement Learning"],
        "Action Verbs": ["Trained", "Developed", "Tested", "Implemented", "Classified"],
        "Quantifiers": ["Millions", "95%", "Hundreds", "20+", "50x"]
    },
    "Business Analysis": {
        "Skills": ["Requirements Gathering", "Process Mapping", "Stakeholder Communication", "Gap Analysis", "KPI Development"],
        "Action Verbs": ["Analyzed", "Identified", "Collaborated", "Documented", "Recommended"],
        "Quantifiers": ["80%", "Dozens", "10x", "Hundreds", "30%"]
    },
    "Finance": {
        "Skills": ["Budgeting", "Forecasting", "Financial Modeling", "Risk Analysis", "Tax Planning", "Portfolio Management"],
        "Action Verbs": ["Prepared", "Evaluated", "Managed", "Optimized", "Advised"],
        "Quantifiers": ["Billions", "95%", "50%", "Thousands", "10+"]
    },
    "Healthcare IT": {
        "Skills": ["EHR", "HIPAA Compliance", "Medical Imaging", "Telemedicine", "HL7", "FHIR"],
        "Action Verbs": ["Implemented", "Developed", "Integrated", "Monitored", "Streamlined"],
        "Quantifiers": ["100%", "Thousands", "50+", "Zero", "10x"]
    },
    "Game Development": {
        "Skills": ["Unity", "Unreal Engine", "C#", "3D Modeling", "Animation", "Game Physics", "Shader Programming"],
        "Action Verbs": ["Created", "Designed", "Optimized", "Rendered", "Published"],
        "Quantifiers": ["Millions", "Thousands", "10x", "5+", "15%"]
    },
    "E-commerce": {
        "Skills": ["Shopify", "Magento", "Inventory Management", "Customer Retention", "Payment Gateways", "Order Fulfillment"],
        "Action Verbs": ["Managed", "Increased", "Automated", "Enhanced", "Designed"],
        "Quantifiers": ["300%", "100+", "Thousands", "Multiple", "10%"]
    },
    "Operations Management": {
        "Skills": ["Six Sigma", "Supply Chain Management", "Logistics", "Lean", "Inventory Optimization", "Vendor Management"],
        "Action Verbs": ["Streamlined", "Reduced", "Improved", "Managed", "Analyzed"],
        "Quantifiers": ["Zero", "Thousands", "20%", "10x", "50+"]
    },
    "DevOps": {
        "Skills": ["Docker", "Kubernetes", "CI/CD", "Jenkins", "Ansible", "Monitoring"],
        "Action Verbs": ["Automated", "Deployed", "Integrated", "Configured", "Provisioned"],
        "Quantifiers": ["95%", "Zero", "Hundreds", "5x", "20+"]
    },
    "Blockchain": {
        "Skills": ["Smart Contracts", "Ethereum", "Bitcoin", "Decentralized Apps", "Solidity", "Consensus Algorithms"],
        "Action Verbs": ["Developed", "Secured", "Verified", "Deployed", "Integrated"],
        "Quantifiers": ["Zero", "Millions", "10x", "Multiple", "50%"]
    },
    "Quality Assurance": {
        "Skills": ["Automation Testing", "Manual Testing", "Bug Tracking", "Performance Testing", "Regression Testing"],
        "Action Verbs": ["Tested", "Debugged", "Validated", "Reported", "Improved"],
        "Quantifiers": ["Thousands", "Zero", "5x", "100+", "15%"]
    },
    "Human Resources": {
        "Skills": ["Recruitment", "Onboarding", "Employee Engagement", "Conflict Resolution", "Payroll Management"],
        "Action Verbs": ["Hired", "Coached", "Evaluated", "Trained", "Facilitated"],
        "Quantifiers": ["500+", "20%", "Multiple", "Hundreds", "10+"]
    },
    "Legal Services": {
        "Skills": ["Contract Drafting", "Legal Research", "Case Management", "Compliance", "Negotiation"],
        "Action Verbs": ["Advised", "Negotiated", "Reviewed", "Drafted", "Represented"],
        "Quantifiers": ["Zero", "Dozens", "100+", "5x", "50%"]
    },
    "Content Writing": {
        "Skills": ["Copywriting", "Editing", "SEO Writing", "Creative Writing", "Proofreading"],
        "Action Verbs": ["Authored", "Edited", "Published", "Researched", "Enhanced"],
        "Quantifiers": ["Thousands", "Hundreds", "Multiple", "10+", "15%"]
    },
}

STRONG_ACTION_VERBS = [
    "Achieved", "Advised", "Analyzed", "Built", "Collaborated", "Conducted", "Created", "Delivered", "Designed", 
    "Developed", "Directed", "Enhanced", "Established", "Executed", "Expanded", "Facilitated", "Generated", 
    "Improved", "Implemented", "Initiated", "Innovated", "Led", "Managed", "Maximized", "Optimized", 
    "Orchestrated", "Planned", "Produced", "Reduced", "Resolved", "Streamlined", "Supervised", "Transformed", 
    "Utilized"
]

QUANTIFIERS = [
    "increased", "reduced", "improved", "generated", "enhanced", "achieved", "delivered", "surpassed", "completed", "exceeded"
]

TRENDING_SKILLS = [
    "AI", "Artificial Intelligence", "Machine Learning", "ML", "AI Models", "Deep Learning", "DL",
    "Data Science", "Data Analytics", "Data Analysis", "Big Data", "Data Engineering", "Data Mining",
    "Natural Language Processing", "NLP", "Computer Vision", "CV", "Python", "Python Programming", "Python Scripting",
    "R Programming", "TensorFlow", "TF", "PyTorch", "Scikit-learn", "Keras", "Deep Neural Networks", "DNN",
    "Artificial Neural Networks", "ANN", "Neural Networks", "Data Visualization", "Tableau", "Power BI",
    "Microsoft Excel", "Excel", "PowerPoint", "PPT", "Word", "Google Analytics", "Google Sheets", "SQL",
    "Structured Query Language", "NoSQL", "MongoDB", "MySQL", "PostgreSQL", "Oracle", "Hadoop", "Spark",
    "AWS", "Amazon Web Services", "Azure", "Google Cloud", "GCP", "Docker", "Kubernetes", "DevOps", "CI/CD",
    "Git", "Version Control", "Blockchain", "Cryptocurrency", "Ethereum", "Smart Contracts", "Dapps",
    "SEO", "Search Engine Optimization", "SEM", "PPC", "Content Marketing", "Marketing Automation",
    "Agile", "Scrum", "Scrum Master", "Project Management", "Jira", "Kanban", "DevOps", "Scrum Framework"
]

BASE_CATEGORY = {
    "Base Category": [
        "sales",
        "ground staff",
        "business development associate",
        "graduate trainee",
        "customer service",
        "client handling",
        "field work",
        "operations",
        "onboarding",
        "trainee",
        "intern",
        "entry level",
        "junior associate",
        "support staff",
        "administration",
        "team assistant",
        "general staff"
    ]
}

# Quality Score Calculation (50% Weightage)
def score_quality(resume_text):
    score = 0

    # Check formatting (e.g., headers, one-page limit)
    headers = ["education", "skills", "experience", "certifications", "summary", "achievements"]
    for header in headers:
        if header in resume_text.lower():
            score += 2  # Assign points for each proper header

    # Check strong action verbs (eliminates duplicates)
    action_verb_score = sum(2 for verb in set(STRONG_ACTION_VERBS) if verb.lower() in resume_text.lower())
    score += min(action_verb_score, 15)  # Add capped action verb score

    # Check quantifiers (eliminates duplicates)
    quantifier_score = sum(2 for quantifier in set(QUANTIFIERS) if quantifier.lower() in resume_text.lower())
    score += min(quantifier_score, 15)  # Add capped quantifier score

    # Check length (favor resumes with 300 to 750 words)
    resume_length = len(resume_text.split())  # Define the resume length
    if 300 <= resume_length <= 750:
        score += 20  # Award more points for resumes of this length
    elif 150 <= resume_length <= 299:
        score += 10
    else:
        score += 5

    return round(min(score, 49), 2)  # Ensure the total score does not exceed 49

# Relevance Score Calculation (45% Weightage)
def score_relevance(resume_text, jd_text):
    # Normalize the texts
    import re
    resume_text = re.sub(r'[^\w\s]', '', resume_text.lower())
    jd_text = re.sub(r'[^\w\s]', '', jd_text.lower())
    
    matching_words = set()
    jd_keywords = set()  # Store keywords found in JD

    # Match keywords in JD
    for keyword, variations in KEYWORD_MAPPINGS.items():
        for variation in variations:
            if variation in jd_text:
                jd_keywords.add(keyword)  # Add to JD keywords
    
    # Match keywords in both JD and resume
    for keyword in jd_keywords:
        for variation in KEYWORD_MAPPINGS[keyword]:
            if variation in resume_text:
                matching_words.add(keyword)
    
    # Calculate relevance score
    if jd_keywords:  # Avoid division by zero
        relevance_score = min(len(matching_words) / len(jd_keywords) * 35, 35)
    elif any(term in jd_text for term in BASE_CATEGORY):  # Check for Base Category terms
        relevance_score = 25  # JD comes under Base Category
    else:
        relevance_score = 0  # No relevant keywords in JD
    
    total_score = 10 + relevance_score  # Add base score of 20
    
    return round(min(total_score, 43), 2)  # Return matching words and capped score


# Trending Skills Score Calculation (5% Weightage)
def score_trending_skills(resume_text):
    # Convert resume text to lowercase for case-insensitive matching
    resume_text = re.sub(r'[^\w\s]', '', resume_text.lower())
    
    # Create a set to track skills found in the resume (avoiding duplicates)
    found_skills = set()

    # Check each skill in the trending skills list
    for skill in TRENDING_SKILLS:
        if skill.lower() in resume_text:
            found_skills.add(skill.lower())  # Add to set to avoid duplicates

    # Calculate the score based on the number of unique skills found
    score = len(found_skills) * 0.2  # Each skill adds 1 to the score
    return round(min(score, 5), 2)  # Cap trending skills score at 5

def show_details(resume_text, jd_text):
    st.write("### 🔍 Detailed Breakdown of Your Resume Score")

    # Display Subcategories in Quality Score Calculation
    st.write("#### 📋 Content Quality Rating - Subcategories:")
    header_score = 0
    action_verb_score = 0
    quantifier_score = 0
    content_score = 0

    # Store matched headers, action verbs, and quantifiers
    matched_headers = set()
    matched_action_verbs = set()
    matched_quantifiers = set()
    
    # Check headers
    headers = ["education", "skills", "experience", "certifications", "summary", "achievements"]
    for header in headers:
        if header in resume_text.lower():
            header_score += 2  # Assign points for each proper header
            matched_headers.add(header)  # Store the matched header

    # Check strong action verbs (eliminates duplicates)
    for verb in set(STRONG_ACTION_VERBS):
        if verb.lower() in resume_text.lower():
            action_verb_score += 2
            matched_action_verbs.add(verb)  # Store the matched verb
    action_verb_score = min(action_verb_score, 15)  # Cap action verb score to a maximum of 15

    # Check quantifiers (eliminates duplicates)
    for quantifier in set(QUANTIFIERS):
        if quantifier.lower() in resume_text.lower():
            quantifier_score += 2
            matched_quantifiers.add(quantifier)  # Store the matched quantifier
    quantifier_score = min(quantifier_score, 15)  # Cap quantifier score to a maximum of 15

    # Check length (favor resumes with 300 to 750 words)
    resume_length = len(resume_text.split())  # Define the resume length
    if 300 <= resume_length <= 750:
        content_score = 20  # Award more points for resumes of this length
    elif 150 <= resume_length <= 299:
        content_score = 10
    else:
        content_score = 5

    # Write the individual scores
    st.write(f"- **Matched Headers:** {', '.join(matched_headers).upper()}")  # Display matched headers
    st.write(f"- **Score for Headers:** {round(header_score, 2)}")
    st.write(f"- **Matched Action Verbs:** {', '.join(matched_action_verbs).upper()}")  # Display matched verbs
    st.write(f"- **Score for Action Verbs:** {round(action_verb_score, 2)}")
    st.write(f"- **Matched Quantifiers:** {', '.join(matched_quantifiers).upper()}")  # Display matched quantifiers
    st.write(f"- **Score for Quantifiers:** {round(quantifier_score, 2)}")
    st.write(f"- **Score for Content (Length):** {round(content_score, 2)}")
    st.write(f"**Total Quality Score:** {min(round(header_score + action_verb_score + quantifier_score + content_score, 2), 49)} / 50")

    # Relevance Score Calculation
    st.write("#### 📊 Job Relevance Assessment - Subcategories:")
    import re

    # Normalize the texts
    resume_text = re.sub(r'[^\w\s]', '', resume_text.lower())
    jd_text = re.sub(r'[^\w\s]', '', jd_text.lower())

    jd_keywords = set()
    matching_keywords = set()

    # Match keywords in JD
    for keyword, variations in KEYWORD_MAPPINGS.items():
        for variation in variations:
            if variation in jd_text:
                jd_keywords.add(keyword)  # Add to JD keywords
    
    # Match keywords in both JD and resume
    for keyword in jd_keywords:
        for variation in KEYWORD_MAPPINGS[keyword]:
            if variation in resume_text:
                matching_keywords.add(keyword)
    
    # Calculate relevance score
    if jd_keywords:  # Avoid division by zero
        relevance_score = min(len(matching_keywords) / len(jd_keywords) * 35, 35)
    elif any(term in jd_text for term in BASE_CATEGORY):  # Check for Base Category terms
        relevance_score = 25  # JD comes under Base Category
    else:
        relevance_score = 0  # No relevant keywords in JD
    
    total_score = 10 + relevance_score  # Add base score of 20

    # Display Sets
    st.write(f"- **Keywords in Job Description:** {', '.join(jd_keywords).upper()}")
    st.write(f"- **Matching Keywords in Resume:** {', '.join(matching_keywords).upper()}")

    total_relevance_score = min(total_score, 43)
    st.write(f"**Total Relevance Score:** {round(total_relevance_score, 2)} / 45")

    st.write("#### 🚀 Emerging Skills Index - Subcategories:")
     # Create a set to track skills found in the resume (avoiding duplicates)
    found_skills = set()

    # Check each skill in the trending skills list
    for skill in TRENDING_SKILLS:
        if skill.lower() in resume_text:
            found_skills.add(skill.lower())  # Add to set to avoid duplicates

    # Calculate the score based on the number of unique skills found
    score = len(found_skills) * 0.2  # Each skill adds 1 to the score
    total_skills_score = min(score, 5)
    st.write(f"- **Emerging Skills in Resume:** {', '.join(found_skills).upper()}")
    st.write(f"**Total Score for Emerging Skills:** {round(total_skills_score, 2)} / 5")
    
# Function: Generate Sample JD 
def generate_sample_files():
    sample_jd = """Software Engineer Job Description:
- Proficient in Python, Java, or similar programming languages
- Hands-on experience with web frameworks like Django or Flask
- Ability to design and maintain APIs
- Familiarity with cloud platforms such as AWS or Azure
- Excellent problem-solving and communication skills"""
    
    # Save sample files as .txt
    with open("sample_jd.txt", "w") as f:
        f.write(sample_jd)

    return "sample_jd.txt"

# Function to create a sample DOCX resume
def create_sample_resume_docx():
    doc = Document()
    doc.add_heading("Sample Resume", level=1)
    doc.add_paragraph("Name: XXX")
    doc.add_paragraph("Email: xxx@example.com")
    doc.add_paragraph("Phone: +XX XXXXXXX")
    doc.add_paragraph("\nSkills:")
    doc.add_paragraph(
        "• Technical Skills: Python | Machine Learning | Data Science | Exploratory Data Analysis (EDA) | Natural Language Processing (NLP)\n"
        "• Tools: Streamlit | Microsoft Office Suite (Excel, Word, PowerPoint) | Pandas | Scikit-Learn | GitHub | Power BI | OpenCV | NLTK\n"
        "• Soft Skills: Problem Solving | Communication | Time Management | Analytical Thinking | Critical Thinking"
    )
    doc.add_paragraph("\nExperience:")
    paragraph = doc.add_paragraph(
        "• Codsoft | Data Science Intern (May 2024)\n"
        "  - Architected dynamic web applications using Python and Streamlit, transformed more than 7 complex datasets into actionable insights and boosted prediction accuracy to over 95% for enhanced decision-making.\n"
        "  - Utilized Seaborn and Matplotlib to create 5 plus interactive dashboards, enhancing data visualization and insights.\n"
    )
    doc.add_paragraph(
        "• PTA, Directorate of Public Instruction (DPI), Government of Tamil Nadu | Data Segmentation Intern (Dec 2023 – Jan 2024)\n"
        "  - Automated and streamlined data segmentation processes using Python and algorithms on over 500,000 continuous records across 50 fields, reduced processing time by over 97%, and saved more than 20 hours per week, enhancing efficiency and accuracy.\n"
        "  - Collaborated with DPI teams to implement data-driven strategies for educational initiatives.\n"
    )
    doc.add_paragraph(
        "• Kindle Direct Publishing | Author (Feb 2023 – Aug 2023)\n"
        "  - Authored and published engaging eBooks on emerging technologies and trending concepts, captivating a diverse audience of tech enthusiasts and professionals, and sold over 55 copies of eBooks.\n"
        "  - Integrated AI-generated insights and data points using NLP techniques to enrich content and provide cutting-edge perspectives, published 4 books, each with over 12 chapters.\n"
    )
    doc.add_paragraph(
        "• Institution of Electronics and Telecommunication Engineers | Data Science Intern (Jun 2023 – Jul 2023)\n"
        "  - Engineered real-time data science applications with dynamic Power BI dashboards and advanced Scikit-Learn models, crafting over 10 visualizations for 4 different applications.\n"
        "  - Formulated and deployed text and image processing techniques using TensorFlow and OpenCV.\n"
    )
    doc.add_paragraph(
        "• YouTube | Content Creator (Jan 2023 – May 2023)\n"
        "  - Created and unveiled engaging content on trending technologies and their practical applications.\n"
        "  - Educated over 350 viewers on Python programming, simplifying complex concepts through clear and concise tutorials.\n"
    )
    doc.add_paragraph("\nEducation:")
    doc.add_paragraph(
        "• SRM Institute of Science and Technology (2022 – 2025)\n"
        "  Bachelor of Computer Applications (BCA) in Data Science | CGPA: 9.79\n"
        "  Relevant Coursework: Machine Learning, Data Science, Natural Language Processing, Artificial Intelligence, Computer Vision, Statistics, Data Engineering, Intelligent Automation, Data Analytics\n"
    )
    doc.add_paragraph("\nProjects:")
    doc.add_paragraph(
        "• Vision Wizard (June 2024)\n"
        "  - Designed Vision Wizard, a drag-and-drop platform reducing pre-processing time by 80%, enabling users with zero programming knowledge to perform more than 12 computer vision tasks using advanced tools and techniques.\n"
    )
    doc.add_paragraph(
        "• TextTrac (May 2024)\n"
        "  - Orchestrated NLP tools to perform 13 plus NLP preprocessing tasks with a single click, reducing normalization time by 95%, and enhancing data-driven decision-making through advanced text manipulation and analysis.\n"
    )
    doc.add_paragraph(
        "• AI Hub (Nov 2023 – Feb 2024)\n"
        "  - Developed AutoDS and AutoNLP solutions, enabling zero-coding, cutting processing time by 70%, and broadening access to over 31 data operations. Revolutionized data processes and democratized advanced analysis and natural language processing.\n"
    )
    doc.add_paragraph("\nCertifications:")
    doc.add_paragraph(
        "• Completed over 30 professional certifications from Coursera, Cisco, NPTEL, Google Skillshop and other platforms, including:\n"
        "  - IBM Data Science Professional Certification • Coursera • August 30, 2023\n"
        "  - The Joy of Computing using Python (IIT Madras) • NPTEL • May 31, 2023\n"
    )
    doc.add_paragraph("\nAchievements:")
    doc.add_paragraph(
        "• Won 2nd prize in the 2024 Project Day for the AI Hub project, hosted by SRM Group.\n"
        "• Received a Silver Medal in the 2024 Research Day for the AI Trinity research work, hosted by SRM IST.\n"
    )
    
    # Save the document to a BytesIO buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Initialize session state
if "resume_file" not in st.session_state:
    st.session_state.resume_file = None

if "jd_text" not in st.session_state: 
    st.session_state.jd_text = None

if "more_details" not in st.session_state:
    st.session_state.more_details = False

# Function to extract text from a file
def extract_text(file):
    try:
        # Handle PDF files
        if file.type == "application/pdf":
            reader = PdfReader(file)
            text = " ".join(page.extract_text() for page in reader.pages if page.extract_text())

        # Handle DOCX files
        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text = docx2txt.process(file)

        # Handle TXT files
        elif file.type == "text/plain":
            text = file.read().decode("utf-8")  # Decode binary content for .txt files

        # Unsupported file type
        else:
            return None
        
        # Return the extracted text after stripping any leading or trailing spaces
        return text.strip() if text else None

    except Exception as e:
        # Log or print the error if needed for debugging
        print(f"Error extracting text: {e}")
        return None

# Shared Function: Clear Inputs
def clear_inputs():
    # Remove variables from session state
    st.session_state.pop("resume_file", None)
    st.session_state.pop("resume_text", None)
    st.session_state.pop("jd_file", None)
    st.session_state.pop("jd_text", None)
    
    # Reset variables to ensure no lingering data
    st.session_state.resume_file = None
    st.session_state.resume_text = None
    st.session_state.jd_file = None
    st.session_state.jd_text = None
    st.session_state.more_details = False

# Function for File Upload Section
def file_upload_section():
    # Resume upload section (only PDF and DOCX allowed)
    st.session_state.resume_file = st.file_uploader("Upload your Resume (PDF, DOCX only):", type=["pdf", "docx"])
    if st.session_state.resume_file:
        st.session_state.resume_text = extract_text(st.session_state.resume_file)
        if not st.session_state.resume_text:
            st.error("Unable to extract text from the resume. Ensure it's a valid and supported file format (PDF, DOCX).")

    # JD input method
    st.radio("How would you like to provide the Job Description?", ["Upload File", "Paste Text"], key="jd_input_method")
    if st.session_state.jd_input_method == "Upload File":
        jd_file = st.file_uploader("Upload Job Description (PDF, DOCX, TXT):", type=["pdf", "docx", "txt"], key="jd_file")
        if jd_file:
            st.session_state.jd_text = extract_text(jd_file)
            if not st.session_state.jd_text:
                st.error("Unable to extract text from the Job Description. Ensure it's a valid and supported file format (PDF, DOCX, TXT).")
    else:
        st.session_state.jd_text = st.text_area("Paste the Job Description:").strip()

# Function: Resume Scoring Logic
def calculate_scores():
    # Ensure text extraction for resume and job description
    resume_text = extract_text(st.session_state.resume_file) if st.session_state.resume_file else st.session_state.resume_text
    if not resume_text:
        st.error("Unable to extract text from the resume. Ensure the format is correct.")
        return None, None, None

    if detect(resume_text) != "en" or detect(st.session_state.jd_text) != "en":
        st.error("Both the resume and job description must be in English.")
        return None, None, None

    quality_score = score_quality(resume_text)
    relevance_score = score_relevance(resume_text, st.session_state.jd_text)
    trending_score = score_trending_skills(resume_text)
    return quality_score, relevance_score, trending_score

# Navigation Menu
st.sidebar.title("📱 Navigation")
page = st.sidebar.radio("**🌐 Select a Feature**", ["Home 🏠", "Simple Resume Score 📝", "Resume Score with Detailed Breakdown 📊", "Recommendations 💡", "Sample JD and Resume 📄"])

# Home Page
if page == "Home 🏠":
    st.header("Welcome to the Resume Score and Recommendation Tool!")
    
    st.write("""
    Maximize your chances of landing the perfect job with our comprehensive resume evaluation tool. 
    Whether you're looking for a simple score, a detailed breakdown of your resume's strengths and weaknesses, 
    or tailored recommendations for improving your resume, we've got you covered!
    """)

    # Simple Resume Score Section
    st.subheader("1. Simple Resume Score 📝")
    st.write("""
    Looking for a quick assessment? Our **Simple Resume Score** feature provides a quick and straightforward evaluation of your resume. 
    All you need to do is upload your resume and job description, and you'll receive a **score out of 100** based on how well your resume aligns with the job you're applying for. 
    This is perfect for those who want a quick overview of their resume quality and job relevance.
    """)

    st.write("🔹 **Features**:")
    st.write("""
    - Provides an overall score of your resume (out of 100).
    - Simple and fast process with no detailed breakdown.
    """)

    # Resume Score with Detailed Breakdown Section
    st.subheader("2. Resume Score with Detailed Breakdown 📊")
    st.write("""
    For a more in-depth analysis of your resume, the **Resume Score with Detailed Breakdown** feature offers an extensive evaluation. 
    This breakdown goes beyond the score and gives you detailed insights into various aspects of your resume. The breakdown includes:
    - **Content Quality Rating**: An evaluation of headers, action verbs, quantifiers, and content length.
    - **Job Relevance Assessment**: How well your resume matches the job description.
    - **Emerging Skills Index**: A measure of how well your resume showcases trending skills in your field.
    """)

    st.write("🔹 **Features**:")
    st.write("""
    - Detailed insights into your resume's strengths and areas for improvement.
    - Breakdown of content quality, job relevance, and trending skills.
    - Actionable feedback for improving your resume.
    """)

    # Recommendations Section
    st.subheader("3. Recommendations 💡")
    st.write("""
    Want to enhance your resume even further? The **Recommendations** feature provides you with personalized recommendations based on the job category you choose. 
    By selecting a job category, you'll receive:
    - **Key Skills**: Essential skills required for the role.
    - **Action Verbs**: Powerful action verbs to use in your resume.
    - **Quantifiers**: Quantitative terms that can make your achievements stand out.
    """)

    st.write("🔹 **Features**:")
    st.write("""
    - Recommendations based on job category.
    - Provides key skills, action verbs, and quantifiers to enhance your resume.
    - Helps you tailor your resume to meet the expectations of the industry.
    """)

    # Encourage user to explore the tool
    st.write("""
    With these three powerful tools, you can take your resume to the next level. Whether you're just starting out or looking to refine your existing resume, 
    we offer the tools you need to stand out to employers!
    """)

    st.write("Get started now by selecting a feature from the navigation bar!")

# Simple Resume Score Page
elif page == "Simple Resume Score 📝":
    st.header("📝 Simple Resume Score")
    file_upload_section()

    if st.button("⚡ Score My Resume"):
        if not st.session_state.resume_file:
            st.error("🚨 Please upload your resume.")
        elif not st.session_state.jd_text:
            st.error("🚨 Please provide the job description.")
        else:
            quality_score, relevance_score, trending_score = calculate_scores()
            if quality_score is not None:
                st.write(f"**Content Quality Rating: {round(quality_score, 2)} / 50**")
                st.write(f"**Job Relevance Assessment: {round(relevance_score, 2)} / 45**")
                st.write(f"**Emerging Skills Index: {trending_score} / 5**")
                final_score = round(quality_score + relevance_score + trending_score, 2)
                st.success(f"🎯 **Your final resume score is: {final_score} / 100**")
                # Provide feedback based on the final score
                if final_score < 70:
                    st.info("🔴 Aim for a score of 70% or higher for better alignment with the job requirements.")
                else:
                    st.success("✅ Great job! Your resume aligns well with the job requirements. Keep it up!")
    
    if st.button("🧹 Clear Inputs"):
        clear_inputs()

# Detailed Breakdown Page
elif page == "Resume Score with Detailed Breakdown 📊":
    st.header("📊 Detailed Resume Breakdown")
    file_upload_section()
    if st.button("⚡ Score My Resume"):
        if not st.session_state.resume_file:
            st.error("🚨 Please upload your resume.")
        elif not st.session_state.jd_text:
            st.error("🚨 Please provide the job description.")
        else:
            # Extract resume text
            resume_text = extract_text(st.session_state.resume_file)
            jd_text = st.session_state.jd_text
            quality_score, relevance_score, trending_score = calculate_scores()
            if quality_score is not None:
                show_details(resume_text, jd_text)
                final_score = round(quality_score + relevance_score + trending_score, 2)
                st.success(f"🎯 **Your final resume score is: {final_score} / 100**")
                # Provide feedback based on the final score
                if final_score < 70:
                    st.info("🔴 Aim for a score of 70% or higher for better alignment with the job requirements.")
                else:
                    st.success("✅ Great job! Your resume aligns well with the job requirements. Keep it up!")
        
    if st.button("🧹 Clear Inputs"):
        clear_inputs()

# Recommendations Page
elif page == "Recommendations 💡":
    st.header("💡 Job Recommendations")
    st.write("Select a job category to get relevant skills, action verbs, and quantifiers to improve your resume.")

    # Job Category Selection
    job_category = st.selectbox("🔍 Choose a Job Category", options=list(CATEGORIES.keys()))

    if st.button("📈 Get Recommendations"):
        if job_category:
            st.subheader(f"🔑 Recommendations for {job_category}")
            
            category_details = CATEGORIES[job_category]
    
            # Display Skills
            st.write("#### 🛠️ Key Skills")
            st.text_area("Skills", ", ".join(category_details["Skills"]), height=100)
    
            # Display Action Verbs
            st.write("#### 💪 Action Verbs")
            st.text_area("Action Verbs", ", ".join(category_details["Action Verbs"]), height=100)
    
            # Display Quantifiers
            st.write("#### 📊 Quantifiers")
            st.text_area("Quantifiers", ", ".join(category_details["Quantifiers"]), height=100)


# JD and Resume Page
elif page == "Sample JD and Resume 📄":
    # Generate and provide download links
    st.header("📄 Sample JD and Resume")

    st.write("""
    Get started by exploring sample Job Descriptions (JD) and Resumes to understand how to tailor your application effectively.
    Use these examples as references to create a compelling resume that aligns with the job you're applying for.
    """)

    st.subheader("Sample Job Description (JD)")
    st.write("""
    **Software Developer Job Description:**
    - Collaborate with cross-functional teams to define, design, and ship new features.
    - Develop, test, and deploy scalable and maintainable software applications.
    - Write clean, efficient, and well-documented code.
    - Participate in code reviews to ensure high-quality standards.
    - Stay updated with the latest industry trends and technologies.
    - Requirements: 
      - Proficiency in programming languages like Python, Java, or JavaScript.
      - Strong problem-solving and analytical skills.
      - Bachelor's degree in Computer Science or related field.
    """)

    # Add download buttons
    jd_file_path = generate_sample_files()
    with open(jd_file_path, "rb") as jd_file:
        st.download_button(
            label="📥 Download Sample JD",
            data=jd_file,
            file_name="sample_jd.txt",
            mime="text/plain",
        )

    st.subheader("Sample Resume")
    st.write("""
    **XXX**  
    Address: XXX, City, State, XXXXX  
    Email: xxx@example.com | Phone: +XX XXXXXXX  

    **Skills:**  
    - **Technical Skills:** Python | Machine Learning | Data Science | Exploratory Data Analysis (EDA) | Natural Language Processing (NLP)  
    - **Tools:** Streamlit | Microsoft Office Suite (Excel, Word, PowerPoint) | Pandas | Scikit-Learn | GitHub | Power BI | OpenCV | NLTK  
    - **Soft Skills:** Problem Solving | Communication | Time Management | Analytical Thinking | Critical Thinking  

    **Experience:**  
    - **Data Science Intern**  
      Codsoft | May 2024  
      Architected dynamic web applications using Python and Streamlit, transformed more than 7 complex datasets into actionable insights and boosted prediction accuracy to over 95% for enhanced decision-making. Utilized Seaborn and Matplotlib to create 5 plus interactive dashboards, enhancing data visualization and insights.  

    - **Data Segmentation Intern**  
      PTA, Directorate of Public Instruction (DPI), Government of Tamil Nadu | Dec 2023 – Jan 2024  
      Automated and streamlined data segmentation processes using Python and algorithms on over 500,000 continuous records across 50 fields, reducing processing time by over 97%, and saving more than 20 hours per week. Collaborated with DPI teams to implement data-driven strategies for educational initiatives.  

    - **Author**  
      Kindle Direct Publishing | Feb 2023 – Aug 2023  
      Authored and published engaging eBooks on emerging technologies and trending concepts, selling over 55 copies of eBooks. Integrated AI-generated insights and data points using NLP techniques to enrich content, publishing 4 books with over 12 chapters each.  

    - **Data Science Intern**  
      Institution of Electronics and Telecommunication Engineers | Jun 2023 – Jul 2023  
      Engineered real-time data science applications with dynamic Power BI dashboards and advanced Scikit-Learn models, crafting over 10 visualizations for 4 different applications. Formulated and deployed text and image processing techniques using TensorFlow and OpenCV.  

    - **Content Creator**  
      YouTube | Jan 2023 – May 2023  
      Created and unveiled engaging content on trending technologies and their practical applications. Educated over 350 viewers on Python programming, simplifying complex concepts through clear and concise tutorials.  

    **Education:**  
    - Bachelor of Computer Applications (BCA) in Data Science  
      SRM Institute of Science and Technology | 2022 – 2025  
      - CGPA: 9.79  
      - Relevant Coursework: Machine Learning, Data Science, Natural Language Processing, Artificial Intelligence, Computer Vision, Statistics, Data Engineering, Intelligent Automation, Data Analytics  

    **Projects:**  
    - **Vision Wizard (June 2024):**  
      Designed Vision Wizard, a drag-and-drop platform reducing pre-processing time by 80%, enabling users with zero programming knowledge to perform more than 12 computer vision tasks using advanced tools and techniques.  

    - **TextTrac (May 2024):**  
      Orchestrated NLP tools to perform 13 plus NLP preprocessing tasks with a single click, reducing normalization time by 95%, and enhancing data-driven decision-making through advanced text manipulation and analysis.  

    - **AI Hub (Nov 2023 – Feb 2024):**  
      Developed AutoDS and AutoNLP solutions, enabling zero-coding, cutting processing time by 70%, and broadening access to over 31 data operations. Revolutionized data processes and democratized advanced analysis and natural language processing.  

    **Certifications:**  
    - Completed over 30 professional certifications from Coursera, Cisco, NPTEL, Google Skillshop, and other platforms, including:  
      - IBM Data Science Professional Certification | Coursera | August 30, 2023  
      - The Joy of Computing using Python (IIT Madras) | NPTEL | May 31, 2023  

    **Achievements:**  
    - Won 2nd prize in the 2024 Project Day for the AI Hub project, hosted by SRM Group.  
    - Received a Silver Medal in the 2024 Research Day for the AI Trinity research work, hosted by SRM IST.  
    """)
   
    # Usage in Streamlit for DOCX
    docx_buffer = create_sample_resume_docx()
    st.download_button(
        label="📥 Download Sample Resume (DOCX)",
        data=docx_buffer,
        file_name="sample_resume.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    
    st.subheader("How to Use:")
    st.write("""
    - Use the sample JD to identify key skills and requirements.
    - Match your resume content to reflect the skills and responsibilities mentioned in the JD.
    - Ensure your resume highlights relevant experience and quantifiable achievements.
    """)
